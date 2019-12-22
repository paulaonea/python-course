import csv
import datetime
from docx import Document


def read_csv(name, attribute):
    with open(name, attribute) as file:
        spreadsheet = csv.DictReader(file)
        for row in spreadsheet:
            date_time.append(datetime.datetime.strptime(row['timestamp'], '%Y-%m-%d %H:%M:%S'))
            bikes.append(int(row['cnt']))
            actual_temperature.append(float(row['t1']))
            feels_like_temperature.append(float(row['t2']))
            humidity.append(float(row['hum']))
            wind_speed.append(float(row['wind_speed']))
            weather_code.append(float(row['weather_code']))
            is_holiday.append(float(row['is_holiday']))
            is_weekend.append(float(row['is_weekend']))
            season.append(float(row['season']))


def statistics(list_of_items, item):
    total = sum(list_of_items)
    average = total / len(list_of_items)
    minimum = min(list_of_items)
    maximum = max(list_of_items)
    document.add_paragraph(f'The average {item} in the data is {average:,.2f}. The {item} ranges from '
                               f'{minimum:,.2f} per month to a maximum of {maximum:,.2f} per month. \n')


date_time = []
bikes = []
temperature = []
actual_temperature = []
feels_like_temperature = []
humidity = []
wind_speed = []
weather_code = []
is_holiday = []
is_weekend = []
season = []

file_name = 'london_merged.csv'

read_csv(file_name, 'r')

document = Document()

document.add_heading(f'Report on London rented bikes', 0)
statistics(bikes, 'number of bikes')
statistics(temperature, 'temperature')

document.save('london_bikes.docx')


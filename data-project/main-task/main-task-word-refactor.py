# compute stats of a simple csv and creates a financial report in word.

import csv
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt


def read_csv(name, attribute):                  # read csv and add each column to a list
    with open(name, attribute) as file:         # all lists are defined as global variables, changed by the function
        spreadsheet = csv.DictReader(file)
        for row in spreadsheet:
            csv_columns[0].append(int(row['sales']))
            csv_columns[1].append(int(row['expenditure']))
            csv_columns[2].append(int(row['sales']) - int(row['expenditure']))
            csv_columns[3].append(row['month'].capitalize())


def statistics(list_of_items):
    total_list = sum(list_of_items)
    avg_list = total_list / len(list_of_items)
    min_list = min(list_of_items)
    max_list = max(list_of_items)
    return total_list, avg_list, min_list, max_list


def print_word_stats(total, average, minimum, maximum, item):
    p = document.add_paragraph(f'Total {item} in the year were ')
    p.add_run(f'{currency}{total: ,.2f}').bold = True
    p.add_run(f', with an average {item} of {currency}{average:,.2f} per month. The {item} range from '
              f'a minimum of {currency}{minimum:,.2f} per month to a maximum of {currency}{maximum:,.2f} per month.')


def charts(list_of_items, item, image_name):
    fig = plt.figure()
    name = fig.add_subplot(1, 1, 1)
    name.bar(csv_columns[3][1:], list_of_items, color='lightblue', edgecolor='darkslategrey')
    name.title = plt.title(f'Total {item} in {currency}', fontsize=18)
    plt.savefig(image_name)
    return


def print_word_chart(image_name):
    document.add_picture(image_name, width=Inches(4.5), height=Inches(3.2))


def profits_details(list_of_items, list_of_months):  # prints profit making and loss making months
    profit_months = ''
    total_profit = 0
    number_of_profit_months = 0
    loss_months = ''
    total_loss = 0
    number_of_loss_months = 0

    for i in range(len(list_of_items)):
        if list_of_items[i] < 0:
            loss_months = f'{loss_months}, {dictionary[list_of_months[i].lower()]}'
            total_loss += list_of_items[i]
            number_of_loss_months += 1
        else:
            profit_months = f'{profit_months}, {dictionary[list_of_months[i].lower()]}'
            total_profit += list_of_items[i]
            number_of_profit_months += 1

    document.add_paragraph(f'There are {number_of_profit_months} profit making months ({profit_months[1:]}) with an '
                           f' average profit of {currency}{total_profit / number_of_profit_months:,.2f} per month.')
    document.add_paragraph(f'There are {number_of_loss_months} loss making months ({loss_months[1:]}) with an average '
                           f'loss of {currency}{total_loss / number_of_loss_months:,.2f} per month. \n')


# ---------------------------

file_name = input('Input the file name?: ')
currency = input('What is the currency?: ')

dictionary = {'jan': 'January', 'feb': 'February', 'mar': 'March', 'apr': 'April', 'may': 'May', 'jun': 'June',
              'jul': 'July', 'aug': 'August', 'sep': 'September', 'oct': 'October', 'nov': 'November',
              'dec': 'December'}

csv_columns = [['sales'], ['costs'], ['profit'], ['months']]
read_csv(file_name, 'r')

document = Document()
document.add_heading(f'Report on {file_name}', 0)

for i in range(3):
    total, average, minimum, maximum = statistics(csv_columns[i][1:])
    charts(csv_columns[i][1:], csv_columns[i][0], f'{csv_columns[i][0]}.png')
    print_word_stats(total, average, minimum, maximum, csv_columns[i][0])
    print_word_chart(f'{csv_columns[i][0]}.png')

profits_details(csv_columns[2][1:], csv_columns[3][1:])

document.save('FinancialReport.docx')

